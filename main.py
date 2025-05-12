import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
from docx2pdf import convert
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

def gerar_documento():
    empresa_nome = entry_empresa_nome.get()
    empresa_razao = entry_empresa_razao.get()
    empresa_endereco = entry_empresa_endereco.get()
    empresa_cnpj = entry_empresa_cnpj.get()
    
    freelancer_nome = entry_freelancer_nome.get()
    freelancer_nacionalidade = entry_nacionalidade.get()
    freelancer_estado_civil = entry_estado_civil.get()
    freelancer_profissao = entry_profissao.get()
    freelancer_cpf = entry_cpf.get()
    freelancer_rg = entry_rg.get()
    freelancer_endereco = entry_endereco.get()

    cidade_forum = entry_forum.get()

    if not all([empresa_nome, empresa_razao, empresa_endereco, empresa_cnpj,
                freelancer_nome, freelancer_nacionalidade, freelancer_estado_civil,
                freelancer_profissao, freelancer_cpf, freelancer_rg, freelancer_endereco,
                cidade_forum]):
        messagebox.showwarning("Campos vazios", "Preencha todos os campos antes de gerar o NDA.")
        return

    doc = Document()
    doc.add_heading('ACORDO DE CONFIDENCIALIDADE (NDA)', 0)

    doc.add_paragraph(f'''PELO PRESENTE INSTRUMENTO PARTICULAR, AS PARTES:

1. CONTRATANTE:
{empresa_nome} ({empresa_razao}), com sede em {empresa_endereco}, inscrita no CNPJ sob o nº {empresa_cnpj}, doravante denominada simplesmente "{empresa_nome}";

2. CONTRATADO(A):
{freelancer_nome}, {freelancer_nacionalidade}, {freelancer_estado_civil}, {freelancer_profissao}, portador(a) do CPF nº {freelancer_cpf} e RG nº {freelancer_rg}, residente e domiciliado(a) à {freelancer_endereco}, doravante denominado(a) simplesmente "FREELANCER";

Têm entre si justo e acordado o presente Acordo de Confidencialidade ("Acordo"), que será regido pelas cláusulas e condições a seguir:
''')

    doc.add_heading('CLÁUSULA 1 – OBJETO', level=1)
    doc.add_paragraph(f'''1.1. O presente Acordo tem por objeto a proteção das informações confidenciais e sensíveis a que o(a) FREELANCER terá acesso no desempenho de suas atividades junto à {empresa_nome}, incluindo, mas não se limitando, aos sistemas Shop9, Cdnet e Safari.''')

    doc.add_heading('CLÁUSULA 2 – DAS INFORMAÇÕES CONFIDENCIAIS', level=1)
    doc.add_paragraph('''2.1. Para fins deste Acordo, consideram-se Informações Confidenciais todas as informações, dados, documentos, códigos, senhas, especificações técnicas, processos, estratégias, bem como qualquer outro material acessado, transmitido ou conhecido pelo(a) FREELANCER no exercício de suas funções, de forma oral, escrita, eletrônica ou qualquer outro meio.

2.2. As Informações Confidenciais incluem, mas não se limitam a:
a) Dados de clientes e fornecedores;
b) Informações operacionais e comerciais;
c) Acesso e uso dos sistemas Shop9, Cdnet e Safari;
d) Credenciais, códigos de acesso e senhas;
e) Estratégias de negócio, planejamento interno e dados financeiros.
''')

    doc.add_heading('CLÁUSULA 3 – OBRIGAÇÕES DO FREELANCER', level=1)
    doc.add_paragraph(f'''3.1. O(a) FREELANCER compromete-se a:
a) Manter sigilo absoluto sobre as Informações Confidenciais;
b) Utilizar tais informações estritamente para o desempenho das atividades acordadas com a {empresa_nome};
c) Não copiar, reproduzir ou divulgar as Informações Confidenciais a terceiros, salvo mediante autorização expressa e por escrito da {empresa_nome};
d) Tomar todas as precauções razoáveis para proteger as Informações Confidenciais de acesso não autorizado.''')

    doc.add_heading('CLÁUSULA 4 – VIGÊNCIA', level=1)
    doc.add_paragraph('''4.1. Este Acordo entra em vigor na data de sua assinatura e permanecerá válido por tempo indeterminado, inclusive após o encerramento da relação entre as partes, até que as informações tornem-se públicas por meios legítimos que não envolvam a violação deste Acordo.''')

    doc.add_heading('CLÁUSULA 5 – PENALIDADES', level=1)
    doc.add_paragraph(f'''5.1. O descumprimento de qualquer obrigação prevista neste Acordo sujeitará o(a) FREELANCER às sanções civis e criminais cabíveis, bem como à indenização por perdas e danos eventualmente causados à {empresa_nome}.''')

    doc.add_heading('CLÁUSULA 6 – DISPOSIÇÕES GERAIS', level=1)
    doc.add_paragraph(f'''6.1. Este Acordo não estabelece vínculo empregatício entre as partes.
6.2. Qualquer alteração ao presente Acordo somente terá validade se formalizada por escrito e assinada por ambas as partes.
6.3. Fica eleito o foro da comarca de {cidade_forum}, com exclusão de qualquer outro, por mais privilegiado que seja, para dirimir quaisquer controvérsias oriundas deste Acordo.''')

    data_hoje = datetime.now().strftime("%d de %B de %Y")
    doc.add_paragraph(f'\n{cidade_forum}, {data_hoje}.\n')
    doc.add_paragraph(f'\n___________________________\n{empresa_nome}\n(nome e assinatura do representante legal)')
    doc.add_paragraph('\n___________________________\nFREELANCER\n(nome e assinatura)')

    docx_file = "Acordo_Confidencialidade_Gerado.docx"


    # Salvar o .docx
    docx_file = "Acordo_Confidencialidade_Gerado.docx"
    doc.save(docx_file)

    # Criar PDF simples com ReportLab
    pdf_file = "Acordo_Confidencialidade_Gerado.pdf"
    c = canvas.Canvas(pdf_file, pagesize=A4)
    width, height = A4

    # Quebrar texto por linhas simples (modo básico, sem formatação complexa)
    texto = [
        f"ACORDO DE CONFIDENCIALIDADE (NDA)\n",
        f"1. CONTRATANTE:\n{empresa_nome} ({empresa_razao}), {empresa_endereco}, CNPJ: {empresa_cnpj}\n",
        f"2. FREELANCER:\n{freelancer_nome}, {freelancer_nacionalidade}, {freelancer_estado_civil}, {freelancer_profissao}, CPF: {freelancer_cpf}, RG: {freelancer_rg}, Endereço: {freelancer_endereco}\n",
        f"OBJETO:\nO presente Acordo tem por objeto a proteção das informações confidenciais relacionadas aos sistemas Shop9, Cdnet e Safari.\n",
        f"INFORMAÇÕES CONFIDENCIAIS:\nDados de clientes, fornecedores, estratégias de negócio e acesso a sistemas.\n",
        f"OBRIGAÇÕES DO FREELANCER:\nManter sigilo, usar informações apenas para as atividades acordadas e proteger contra acesso não autorizado.\n",
        f"VIGÊNCIA:\nEste acordo é válido por tempo indeterminado.\n",
        f"PENALIDADES:\nO descumprimento acarretará sanções civis e criminais.\n",
        f"DISPOSIÇÕES GERAIS:\nForo: {cidade_forum}\n",
        f"\n{cidade_forum}, {data_hoje}\n\n___________________________\n{empresa_nome}\n\n___________________________\nFREELANCER"
    ]

    # Adicionar ao PDF
    y = height - 40
    for linha in texto:
        for sub in linha.split("\n"):
            c.drawString(40, y, sub)
            y -= 18
            if y < 50:
                c.showPage()
                y = height - 40

    c.save()

    messagebox.showinfo("Sucesso", "Documento Word e PDF gerados com sucesso!")


# GUI com Estilo
root = tk.Tk()
root.title("Gerador de NDA")
root.geometry("600x750")
root.resizable(False, False)

style = ttk.Style()
style.configure("TLabel", font=("Segoe UI", 10))
style.configure("TEntry", font=("Segoe UI", 10))
style.configure("TButton", font=("Segoe UI", 11, "bold"), padding=6)

def criar_campo(parent, texto):
    ttk.Label(parent, text=texto).pack(anchor="w", padx=10, pady=2)
    campo = ttk.Entry(parent, width=60)
    campo.pack(padx=10, pady=2)
    return campo

frame_contratante = ttk.LabelFrame(root, text="Dados do CONTRATANTE")
frame_contratante.pack(fill="x", padx=20, pady=10)

entry_empresa_nome = criar_campo(frame_contratante, "Nome da Empresa")
entry_empresa_razao = criar_campo(frame_contratante, "Razão Social")
entry_empresa_endereco = criar_campo(frame_contratante, "Endereço Completo")
entry_empresa_cnpj = criar_campo(frame_contratante, "CNPJ")

frame_freelancer = ttk.LabelFrame(root, text="Dados do FREELANCER")
frame_freelancer.pack(fill="x", padx=20, pady=10)

entry_freelancer_nome = criar_campo(frame_freelancer, "Nome Completo")
entry_nacionalidade = criar_campo(frame_freelancer, "Nacionalidade")
entry_estado_civil = criar_campo(frame_freelancer, "Estado Civil")
entry_profissao = criar_campo(frame_freelancer, "Profissão")
entry_cpf = criar_campo(frame_freelancer, "CPF")
entry_rg = criar_campo(frame_freelancer, "RG")
entry_endereco = criar_campo(frame_freelancer, "Endereço Completo")
entry_forum = criar_campo(frame_freelancer, "Cidade/UF do Foro")

ttk.Button(root, text="Gerar NDA", command=gerar_documento).pack(pady=20)

root.mainloop()
